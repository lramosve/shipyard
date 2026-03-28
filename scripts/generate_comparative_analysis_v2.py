"""Generate the Comparative Analysis V2 document as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ── Helper functions ────────────────────────────────────────────────────
def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


# ── Title Page ──────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Comparative Analysis — V2")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "SuperShip (Agent-Built) vs. FleetGraph (Original)\n"
    "Shipyard Autonomous Coding Agent Project\n"
    "March 2026"
)
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ── 1. Executive Summary ────────────────────────────────────────────────
doc.add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "Shipyard is an autonomous coding agent built with LangGraph, FastAPI, and GPT-4o "
    "that was directed to rebuild the Ship application—a government-grade project "
    "management platform—from scratch. The original Ship codebase (FleetGraph) is a "
    "mature TypeScript monorepo spanning approximately 34,800 lines of production code "
    "across 447 source files, with 609 commits, 71 Playwright end-to-end test suites, "
    "PostgreSQL with 30+ migrations, real-time collaborative editing via TipTap/Yjs, "
    "a LangGraph-powered proactive intelligence agent (FleetGraph), PIV/OIDC authentication, "
    "and WCAG 2.1 AA accessibility compliance."
)

doc.add_paragraph(
    "The agent-built version (SuperShip) is a pnpm monorepo totaling approximately 1,711 "
    "lines of TypeScript across three packages (shared schemas, Express.js API, React frontend). "
    "It was produced in 6 commits through 8 development phases, implementing a unified document "
    "model with Zod-validated schemas, in-memory repositories, a document CRUD dashboard, "
    "a parity console (findings, trace events, chat), and deployment to Vercel (frontend) "
    "and Railway (API). SuperShip faithfully preserves the original's core architectural "
    "principle—'everything is a document'—but operates entirely in-memory without persistence, "
    "authentication, or real-time collaboration."
)

doc.add_paragraph(
    "This V2 analysis reflects a second rebuild pass that targeted architectural fidelity "
    "over feature breadth. Where the first rebuild (new_ship, Python/FastAPI, ~5,500 LOC) "
    "swapped the language stack and adopted conventional normalized tables, SuperShip preserves "
    "the original's TypeScript/Express/React stack, its unified document model, its Zod schema "
    "contracts, and its monorepo workspace structure. The trade-off: SuperShip implements fewer "
    "features than the first rebuild (no Kanban, no Gantt, no analytics) but achieves "
    "substantially higher architectural fidelity to the original."
)

# ── 2. Architectural Comparison ─────────────────────────────────────────
doc.add_heading("2. Architectural Comparison", level=1)

doc.add_heading("2.1 Technology Stack", level=2)

doc.add_paragraph(
    "Unlike the first rebuild, which swapped TypeScript for Python, SuperShip preserves "
    "the original's technology choices across every layer:"
)

add_table(
    ["Layer", "FleetGraph (Original)", "SuperShip (Agent-Built)"],
    [
        ["Runtime", "Node.js >= 20, pnpm >= 9", "Node.js >= 20, pnpm >= 9"],
        ["Language", "TypeScript 5.7 (strict)", "TypeScript 5.7 (strict)"],
        ["Backend", "Express.js 4.21", "Express.js 4.21"],
        ["Frontend", "React 18 + Vite 6 + TipTap/Yjs", "React 18 + Vite 6 (no editor)"],
        ["Validation", "Zod 3.24 + OpenAPI", "Zod 3.24"],
        ["Database", "PostgreSQL 16 (raw SQL, 30+ migrations)", "In-memory (no persistence)"],
        ["State Mgmt", "TanStack Query + React Context", "useState/useEffect hooks"],
        ["Real-time", "Yjs WebSocket CRDT sync", "None"],
        ["AI Agent", "FleetGraph (LangGraph JS, Claude Sonnet 4)", "None"],
        ["Auth", "Session + PIV/OIDC + bcrypt + CSRF + rate-limit", "None"],
        ["CSS", "TailwindCSS 3 + USWDS", "Inline CSS-in-JS"],
        ["Testing", "Vitest + Playwright (71 E2E suites)", "Vitest + Supertest (19 test cases)"],
        ["Deployment", "Vercel + Railway + Terraform (AWS GovCloud)", "Vercel + Railway"],
        ["Monorepo", "pnpm workspaces (api, web, shared)", "pnpm workspaces (api, web, shared)"],
    ],
)

doc.add_heading("2.2 Data Model", level=2)

doc.add_paragraph(
    "The most significant architectural achievement of the SuperShip rebuild is its faithful "
    "adoption of the unified document model. The original FleetGraph stores every entity—wiki "
    "pages, issues, programs, projects, sprints, persons, weekly plans, retros, standups, "
    "and reviews—as rows in a single 'documents' table, discriminated by document_type, with "
    "type-specific data in JSONB 'properties' columns. Relationships are managed through a "
    "separate 'document_associations' junction table."
)

doc.add_paragraph(
    "SuperShip reproduces this pattern faithfully. All 9 document types (wiki, issue, program, "
    "project, week, weekly_plan, weekly_retro, person, view) are defined in a shared Zod schema "
    "package. The InMemoryDocumentRepository stores documents as a flat array discriminated by "
    "type, with status (draft, active, archived), tags, and owner metadata. This is a marked "
    "improvement over the first rebuild (new_ship), which used six separate normalized tables "
    "and lost the architectural insight that makes the original's data model powerful."
)

doc.add_paragraph(
    "However, SuperShip's document model lacks the depth of the original's:"
)

add_table(
    ["Feature", "FleetGraph", "SuperShip"],
    [
        ["Document types", "10 (incl. standup, weekly_review)", "9 (no standup or weekly_review)"],
        ["Type-specific properties", "Rich JSONB (IssueProperties, ProgramProperties, etc.)", "None—all docs share same flat fields"],
        ["Relationships", "document_associations table + document_links (backlinks)", "None"],
        ["Hierarchy", "parent_id with drag-and-drop reordering", "None"],
        ["Versioning", "document_history table + Yjs state", "None"],
        ["Conversion", "Document type conversion with snapshots", "None"],
        ["Soft delete", "deleted_at with 30-day retention", "None"],
        ["Visibility", "private/workspace with RBAC enforcement", "None"],
    ],
)

doc.add_heading("2.3 Parity Artifacts", level=2)

doc.add_paragraph(
    "SuperShip introduces a 'parity console' layer that does not exist in the original FleetGraph. "
    "This includes three artifact types: Findings (proactive quality/risk signals with severity "
    "and status), Trace Events (audit trail: document_created, document_updated, finding_detected, "
    "chat_generated, status_changed), and Chat Messages (per-document asynchronous conversations). "
    "While FleetGraph has a more sophisticated equivalent—the FleetGraph agent with proactive polling, "
    "LLM-powered analysis, and human-in-the-loop approval—SuperShip's parity layer represents "
    "a deterministic scaffolding that could be wired to an actual LLM in a future phase."
)

doc.add_heading("2.4 Choices the Agent Made That a Human Would Not", level=2)

bullets = [
    "In-memory storage with no persistence layer. A human developer would reach for PostgreSQL "
    "(or at minimum SQLite) immediately, especially given that the original uses PostgreSQL with "
    "30+ migrations. The agent optimized for speed-to-demo over operational durability.",

    "All frontend code in a single 772-line main.tsx file with 11 useState hooks. The original "
    "separates concerns across 45+ components, 20+ hooks, 8+ context providers, and dedicated "
    "page modules. A human would structure the component hierarchy from the start.",

    "Inline CSS-in-JS for all styling instead of TailwindCSS or a component library. The original "
    "uses Tailwind with USWDS compliance. A human would adopt Tailwind early to avoid "
    "accumulating technical debt in inline styles.",

    "Deterministic chat responses (template strings) instead of actual LLM integration. The "
    "parity console's chat feature always returns the same assistant response regardless of "
    "user input. A human would at minimum stub a real API call.",

    "No authentication or authorization whatsoever. Every endpoint is public. The original has "
    "session-based auth with PIV/OIDC, CSRF protection, rate limiting, and per-workspace RBAC. "
    "A human would add at least basic auth before deploying.",

    "Seeded data hardcoded in the shared schema package rather than a database seed script or "
    "migration. This conflates schema definitions with test data.",
]

for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

# ── 3. Performance Benchmarks ───────────────────────────────────────────
doc.add_heading("3. Performance Benchmarks", level=1)

doc.add_heading("3.1 Lines of Code", level=2)

add_table(
    ["Metric", "FleetGraph", "SuperShip", "Ratio"],
    [
        ["Total production LOC", "~34,800", "~1,711", "20:1"],
        ["Backend LOC", "~15,000+ (138 files)", "269 (1 file)", "56:1"],
        ["Frontend LOC", "~14,000+ (205+ files)", "772 (1 file)", "18:1"],
        ["Shared types LOC", "~1,500 (8 files)", "289 (1 file)", "5:1"],
        ["Test LOC", "~200,000+ (105 test files)", "355 (3 files)", "563:1"],
        ["Database schema LOC", "~20,700 (schema.sql)", "0", "N/A"],
        ["Total commits", "609", "6", "102:1"],
    ],
)

doc.add_paragraph(
    "The 20:1 ratio in production code reflects both genuine simplification (in-memory storage "
    "eliminates all database code) and significant feature omission. The 563:1 ratio in test "
    "code is the starkest difference—FleetGraph's 71 Playwright E2E suites and 34 API unit test "
    "files represent months of accumulated quality assurance that SuperShip's 19 test cases "
    "cannot replicate."
)

doc.add_heading("3.2 Feature Coverage", level=2)

add_table(
    ["Feature Category", "FleetGraph", "SuperShip", "Coverage"],
    [
        ["Document CRUD", "Full (10 types, JSONB properties)", "Full (9 types, flat fields)", "Partial"],
        ["Document relationships", "Associations, backlinks, hierarchy", "None", "Missing"],
        ["Issue tracking", "State machine, priority, kanban, assignees", "None (issue is a doc type only)", "Missing"],
        ["Weekly workflows", "Plans, retros, standups, approvals", "None (plan/retro are doc types only)", "Missing"],
        ["Team management", "Directory, capacity, org chart, RACI", "None (person is a doc type only)", "Missing"],
        ["Real-time collaboration", "TipTap + Yjs CRDT + WebSocket", "None", "Missing"],
        ["AI intelligence agent", "FleetGraph (proactive + on-demand)", "Deterministic parity console", "Scaffolded"],
        ["Authentication", "Session + PIV/OIDC + API tokens", "None", "Missing"],
        ["Authorization", "Workspace RBAC + visibility", "None", "Missing"],
        ["Audit logging", "Full (audit_logs table, compliance-grade)", "Trace events (seeded, static)", "Scaffolded"],
        ["Search", "Full-text (title + content)", "Client-side filter (type + text)", "Partial"],
        ["File uploads", "S3 + CloudFront CDN", "None", "Missing"],
        ["Accessibility", "WCAG 2.1 AA + Section 508 + USWDS", "None", "Missing"],
        ["API documentation", "OpenAPI/Swagger (auto-generated)", "None", "Missing"],
        ["Database persistence", "PostgreSQL (30+ migrations)", "In-memory (data lost on restart)", "Missing"],
        ["Deployment infra", "Terraform (AWS GovCloud)", "Vercel + Railway (manual)", "Partial"],
        ["CI/CD", "GitHub Actions (build + test + deploy)", "None", "Missing"],
        ["Comments/threading", "Inline TipTap comments with replies", "None", "Missing"],
        ["Command palette", "CMD+K navigation", "None", "Missing"],
    ],
)

doc.add_heading("3.3 Dependency Footprint", level=2)

add_table(
    ["Metric", "FleetGraph", "SuperShip"],
    [
        ["Production deps (API)", "~25 packages", "4 packages"],
        ["Production deps (Web)", "~30 packages", "4 packages"],
        ["Dev deps (total)", "~40 packages", "~15 packages"],
        ["Lock file size", "365 KB", "96 KB"],
        ["Critical deps", "Express, React, pg, Yjs, TipTap, TanStack Query, Zod, Tailwind", "Express, React, Zod, Vite"],
    ],
)

doc.add_paragraph(
    "SuperShip's minimal dependency footprint is both a strength (smaller attack surface, "
    "faster installs, fewer version conflicts) and a reflection of missing functionality. "
    "The absence of pg, Yjs, TipTap, TanStack Query, Tailwind, and helmet represents features "
    "not yet implemented rather than deliberate simplification."
)

doc.add_heading("3.4 Build and Test Performance", level=2)

add_table(
    ["Metric", "FleetGraph", "SuperShip"],
    [
        ["Type-check time", "~15-20s (447 files)", "~3-5s (6 files)"],
        ["Build time", "~30-45s (full workspace)", "~5-10s (full workspace)"],
        ["Unit test count", "~100+ API tests + 4 frontend tests", "14 API tests + 5 frontend tests"],
        ["E2E test count", "71 Playwright suites", "0"],
        ["Release verification", "Manual + CI pipeline", "Automated script (verify-release.mjs)"],
    ],
)

# ── 4. Shortcomings ─────────────────────────────────────────────────────
doc.add_heading("4. Shortcomings", level=1)

doc.add_paragraph(
    "This section catalogs every significant shortcoming identified during and after the "
    "SuperShip rebuild, with evidence from the rebuild log and post-analysis comparison."
)

doc.add_heading("4.1 Interventions During Rebuild", level=2)

add_table(
    ["#", "What Broke / Got Stuck", "What Was Done", "What It Reveals"],
    [
        ["1", "Agent produced all frontend code in a single main.tsx",
         "Accepted—refactoring mid-stream risked introducing bugs",
         "File-editing agents gravitate toward fewer, larger files because splitting requires understanding import graphs and shared state. The agent lacks a 'code architecture' sense."],
        ["2", "In-memory repositories lose all data on restart",
         "Accepted for MVP scope—documented as Phase 6+ roadmap item",
         "The agent optimizes for 'does it run?' not 'is it durable?' Without explicit instruction to add persistence, it chose the path of least resistance."],
        ["3", "Parity chat returns deterministic responses",
         "Accepted—wiring to real LLM was out of scope",
         "The agent scaffolds interfaces before implementations. The chat endpoint exists with correct schemas but no intelligence behind it."],
        ["4", "No test coverage for error paths or edge cases",
         "Not intervened—19 happy-path tests accepted for deadline",
         "The agent writes tests that validate the golden path. Negative testing (invalid input, concurrent access, network failure) requires explicit prompting."],
        ["5", "Seed data mixed into shared schema package",
         "Not intervened—accepted as architectural debt",
         "The agent doesn't distinguish between schema definitions and test fixtures. It puts data wherever is most convenient for the current task."],
        ["6", "No TypeScript path aliases or import organization",
         "Not intervened—flat structure manageable at current scale",
         "The agent doesn't anticipate growth. Import organization matters at 50+ files but is invisible at 6 files."],
    ],
)

doc.add_heading("4.2 Feature Omissions", level=2)

doc.add_paragraph(
    "The following features from FleetGraph were not rebuilt in SuperShip. Each omission "
    "is categorized by root cause."
)

p("Omissions due to complexity (agent could not have built without extensive guidance):", bold=True)
bullets = [
    "FleetGraph intelligence agent — requires understanding of LangGraph graph design, "
    "conditional fan-out/fan-in, proactive polling schedules, and human-in-the-loop approval "
    "workflows. Building an AI agent inside an application is a meta-task that the coding agent "
    "cannot infer from API surface alone.",
    "Real-time collaborative editing — TipTap + Yjs CRDT integration involves WebSocket "
    "protocol handling, binary state serialization, conflict resolution, offline tolerance "
    "with IndexedDB, and debounced persistence. This is irreducibly complex.",
    "PIV/OIDC authentication — government PKI integration requires understanding of X.509 "
    "certificates, FPKI validators, CAIA OAuth flows, and certificate-to-user binding.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

p("Omissions due to scope (agent could build with sufficient time and prompting):", bold=True)
bullets = [
    "PostgreSQL persistence — straightforward migration from in-memory to pg driver with SQL",
    "Session-based authentication with bcrypt — standard Express.js middleware pattern",
    "RBAC and workspace isolation — query filtering by workspace_id, role checks in middleware",
    "TailwindCSS styling — class-based replacement of inline styles",
    "Full-text search — PostgreSQL tsvector or simple LIKE queries",
    "File uploads — multer middleware + S3 client",
    "Audit logging — INSERT into audit_logs on each mutation",
    "CI/CD pipeline — GitHub Actions YAML generation",
    "OpenAPI documentation — Zod-to-OpenAPI is already a dep in the original",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("4.3 Quality Gaps", level=2)

add_table(
    ["Quality Dimension", "FleetGraph", "SuperShip", "Gap Severity"],
    [
        ["Type safety", "Strict TS + Zod + discriminated unions", "Strict TS + Zod", "Low"],
        ["Error handling (API)", "Structured ApiResponse<T> union, HTTP codes, Zod errors", "400/404 with Zod flatten", "Low"],
        ["Error handling (UI)", "React error boundaries + TanStack retry", "Try-catch with fallback state", "Medium"],
        ["Security headers", "Helmet + CSRF + rate limiting", "None", "Critical"],
        ["Input sanitization", "TipTap + DOMPurify", "Zod validation only", "Medium"],
        ["Accessibility", "WCAG 2.1 AA, axe-core E2E tests, USWDS", "No accessibility features", "Critical"],
        ["Performance optimization", "TanStack Query caching, useMemo, pagination", "useMemo for filtering only", "Medium"],
        ["Code organization", "45+ components, 20+ hooks, 8+ contexts", "1 monolithic component file", "High"],
        ["Documentation", "47 markdown files, JSDoc, architecture docs", "README + architecture.md", "Medium"],
    ],
)

# ── 5. Advances ─────────────────────────────────────────────────────────
doc.add_heading("5. Advances", level=1)

doc.add_paragraph(
    "Despite the significant feature gap, the SuperShip rebuild demonstrates several areas "
    "where the autonomous agent outperformed or moved faster than manual development."
)

doc.add_heading("5.1 Architectural Fidelity to the Unified Document Model", level=2)
doc.add_paragraph(
    "The most notable advance over the first rebuild is that SuperShip correctly identifies "
    "and preserves the original's core architectural principle: everything is a document. "
    "The first rebuild (new_ship) missed this entirely, creating six normalized tables. "
    "SuperShip's shared Zod schema package defines all 9 document types with a discriminated "
    "DocumentType enum, mirroring FleetGraph's approach. This demonstrates that with better "
    "context injection—specifically, feeding the agent the original's architecture documentation "
    "before starting—the agent can make architecturally faithful decisions."
)

doc.add_heading("5.2 Speed of Delivery", level=2)
doc.add_paragraph(
    "SuperShip was produced in 6 commits across 8 development phases. The agent generated "
    "a working, deployed, tested application with shared schema contracts, a REST API, a "
    "multi-view React frontend, and a parity console in a fraction of the time a human "
    "developer would need. The entire shared types package (289 LOC of Zod schemas with 15 "
    "schema definitions and 17 derived TypeScript types) was produced in a single pass."
)

doc.add_heading("5.3 Schema-First Development", level=2)
doc.add_paragraph(
    "The agent adopted a schema-first approach: defining all types in the shared package "
    "before building the API or frontend. This mirrors best practices and ensures type safety "
    "across the full stack. Both the API's input validation and the frontend's form handling "
    "reference the same Zod schemas, eliminating the class of bugs where API and client "
    "disagree on field names or types."
)

doc.add_heading("5.4 Deployment Pipeline", level=2)
doc.add_paragraph(
    "The agent produced deployment configurations for both Vercel (frontend) and Railway (API) "
    "including a nixpacks.toml for Railway's build system, environment variable configuration, "
    "and a health check endpoint. It also generated a release verification script "
    "(verify-release.mjs) that automates end-to-end testing of the deployed API—starting the "
    "server, polling for readiness, running CRUD + parity operations, and asserting response "
    "structures. This level of release automation was not present in the first rebuild."
)

doc.add_heading("5.5 Parity Console as Observability Scaffolding", level=2)
doc.add_paragraph(
    "While deterministic, the parity console (findings, trace events, chat) establishes the "
    "API contracts and UI patterns needed for a real observability layer. The trace event "
    "taxonomy (document_created, document_updated, finding_detected, chat_generated, "
    "status_changed) mirrors the original's audit_logs table in intent, and the findings "
    "model (severity + status) parallels FleetGraph's proactive detection output. Replacing "
    "the deterministic responses with LLM calls would require changing only the repository "
    "implementation, not the API or frontend code."
)

# ── 6. Trade-off Analysis ───────────────────────────────────────────────
doc.add_heading("6. Trade-off Analysis", level=1)

doc.add_paragraph(
    "This section evaluates the major architecture decisions in the Shipyard agent and "
    "whether they were the right call, informed by both rebuilds."
)

doc.add_heading("6.1 Anchor-Based Editing (Exact String Match)", level=2)
p("Verdict: Correct for this project scope.", bold=True)
doc.add_paragraph(
    "The agent's edit_file tool uses exact string matching with uniqueness enforcement—the "
    "same strategy as Claude Code and aider. For files under 300 lines (SuperShip's largest "
    "file is 772 lines), this worked reliably. The triple validation (file-must-be-read, "
    "mtime-must-match, exactly-one-match) prevented the most common failure modes. The "
    "fuzzy-match hint for near-misses and line-number reporting for ambiguous matches enabled "
    "self-correction. If the project required editing files over 1,000 lines, AST-based editing "
    "would be worth the per-language parser cost. For this scope, anchor-based was the right call."
)

doc.add_heading("6.2 LangGraph over Custom Loop", level=2)
p("Verdict: Correct, with caveats.", bold=True)
doc.add_paragraph(
    "LangGraph provided automatic state management, conditional routing, and LangSmith tracing "
    "with minimal code. The add_messages reducer correctly handled message list merging, and "
    "bind_tools auto-generated tool schemas. The trade-off is framework coupling—if LangGraph "
    "changes its API, migration is non-trivial. For a one-week sprint, the productivity gain "
    "justified the coupling. For a production agent, the custom loop would provide more control "
    "over retry logic, streaming, and error recovery."
)

doc.add_heading("6.3 GPT-4o over Claude for the Rebuild", level=2)
p("Verdict: Pragmatically correct, architecturally suboptimal.", bold=True)
doc.add_paragraph(
    "GPT-4o was chosen because the developer had more available tokens on the OpenAI API. "
    "This was the right pragmatic call—token budget was the binding constraint. However, "
    "GPT-4o showed specific weaknesses: more verbose tool call arguments, occasional empty "
    "content fields requiring extraction workarounds, and (notably) the first rebuild's "
    "language swap to Python, which suggests GPT-4o defaults to Python more aggressively "
    "than Claude does. The SuperShip rebuild (which preserved TypeScript) benefited from "
    "better context injection, not a model change."
)

doc.add_heading("6.4 In-Memory Storage in SuperShip", level=2)
p("Verdict: Wrong for anything beyond a demo.", bold=True)
doc.add_paragraph(
    "The decision to use in-memory repositories means all data is lost on server restart. "
    "This was expedient—it eliminated database setup, migrations, and connection management—but "
    "it makes SuperShip fundamentally non-functional as a real application. The original uses "
    "PostgreSQL with 30+ carefully designed migrations and a 20,700-line schema. Even SQLite "
    "(which the first rebuild used) would have been better. This is the single highest-impact "
    "improvement for a hypothetical Phase 2."
)

doc.add_heading("6.5 Supervisor Multi-Agent Pattern", level=2)
p("Verdict: Correct design, underutilized in practice.", bold=True)
doc.add_paragraph(
    "The supervisor/coder/researcher pattern cleanly separates read-only investigation from "
    "write operations, prevents context bleed between workers, and enables parallel task "
    "execution. In practice, most SuperShip rebuild tasks were dispatched to the coder alone "
    "because the tasks were predominantly generative (write new code) rather than investigative "
    "(understand existing code). The researcher worker would prove more valuable on maintenance "
    "tasks—debugging, refactoring, or extending an existing codebase—where understanding the "
    "current state is as important as generating new code."
)

doc.add_heading("6.6 Context Compaction at 50% Window", level=2)
p("Verdict: Slightly aggressive but defensible.", bold=True)
doc.add_paragraph(
    "Compacting at 50% of the 200K context window triggers earlier than necessary, potentially "
    "losing useful context. But the safety margin proved valuable: a single read_file on a large "
    "file can add 20K+ characters, and the orphaned-ToolMessage bug showed that compaction boundary "
    "handling is error-prone. The 50% threshold prevented out-of-context failures during the "
    "SuperShip rebuild. A 60-65% threshold would be the optimal balance for future iterations."
)

# ── 7. If You Built It Again ────────────────────────────────────────────
doc.add_heading("7. If You Built It Again", level=1)

doc.add_heading("7.1 Architecture Changes", level=2)

bullets = [
    "Add a planning phase before code generation. The agent should read the target codebase's "
    "architecture documentation, README, and schema before writing any code. The SuperShip "
    "rebuild's architectural fidelity came from injecting context about the unified document "
    "model—the first rebuild (new_ship) failed precisely because this context was absent. "
    "A dedicated 'architecture analysis' step that produces a design document before the first "
    "edit would dramatically improve output quality.",

    "Implement a file organization agent. The current agent creates files where convenient, "
    "leading to monolithic components. A specialized agent that periodically reviews file sizes "
    "and proposes splits—or that enforces a max-lines-per-file rule—would produce better-structured "
    "codebases.",

    "Add a security review pass. Neither rebuild caught security issues (SHA-256 in new_ship, "
    "no auth in SuperShip). A post-generation security audit tool that checks for common "
    "vulnerabilities (OWASP Top 10, credential handling, CORS configuration) and generates "
    "fix tasks would address this systematic blind spot.",

    "Implement incremental feature delivery. Instead of generating all code in one pass, "
    "the agent should build feature-by-feature with tests after each feature. This matches "
    "how SuperShip was built (8 phases) and produced better results than the first rebuild's "
    "more monolithic approach.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("7.2 File Editing Strategy Changes", level=2)

bullets = [
    "Add a 'refactor' tool that splits large files. The current toolset (read_file, edit_file, "
    "write_file) operates on individual files but has no concept of extracting a component from "
    "one file into a new file while updating imports. This is the most common operation a human "
    "developer performs on agent-generated code.",

    "Implement multi-file atomic edits. Some changes require modifying multiple files "
    "simultaneously (e.g., renaming a type in shared/ requires updating api/ and web/). The "
    "current sequential edit model can leave the codebase in an inconsistent state mid-operation.",

    "Add AST-aware validation for TypeScript/JavaScript. After each edit, run a quick type-check "
    "(tsc --noEmit) to catch syntax errors immediately rather than discovering them when the user "
    "tries to build. The current approach relies on the agent noticing errors in subsequent tool "
    "calls.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("7.3 Context Management Changes", level=2)

bullets = [
    "Implement semantic context injection. Instead of dumping entire files into context, extract "
    "and inject only the relevant sections (function signatures, type definitions, API contracts). "
    "This would allow the agent to work with larger codebases without hitting context limits.",

    "Add a 'codebase map' that persists across sessions. A lightweight index of file paths, "
    "exported symbols, and dependency relationships would let the agent navigate large codebases "
    "without reading every file. The current approach (search_files + list_files) requires "
    "re-discovering the codebase structure on every session.",

    "Improve compaction to preserve architectural decisions. The current summarizer treats all "
    "messages equally. Messages containing architectural decisions (e.g., 'use the unified "
    "document model') should be preserved at higher priority than routine tool call results.",

    "Add explicit 'design constraint' context that survives compaction. Architectural requirements "
    "like 'preserve the TypeScript stack' or 'use the unified document model' should be stored "
    "as persistent system instructions, not as conversation messages that can be compacted away.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("7.4 What Would Be Different Overall", level=2)

doc.add_paragraph(
    "The fundamental lesson from both rebuilds is that autonomous coding agents are excellent "
    "code generators but poor architects. They produce correct, well-typed, working code at "
    "remarkable speed—but they make architectural decisions based on what is easiest to generate, "
    "not what best serves the application's long-term needs. The first rebuild swapped languages "
    "and flattened the data model. The second rebuild preserved the architecture but dropped "
    "most features."
)

doc.add_paragraph(
    "If built again, the agent would separate planning from execution entirely. A planning "
    "phase would analyze the target, produce an architectural specification, and define feature "
    "priorities. An execution phase would implement features incrementally against that spec, "
    "with automated quality gates (type-check, lint, test, security scan) after each feature. "
    "A review phase would compare output against the original, flag gaps, and iterate. This "
    "three-phase approach—plan, execute, review—mirrors how experienced human teams work and "
    "would produce substantially better results than either single-pass rebuild achieved."
)

# ── Save ────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "..", "docs", "Comparative Analysis - V2.docx")
output_path = os.path.normpath(output_path)
doc.save(output_path)
print(f"Saved to {output_path}")
